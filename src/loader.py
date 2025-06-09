from llama_index.core import Document
import fitz  # PyMuPDF
import docx
from io import BytesIO
import streamlit as st

def load_documents_from_upload(uploaded_files):
    docs = []

    for uploaded_file in uploaded_files:
        file_content = uploaded_file.read()

        try:
            if uploaded_file.type == "text/plain":
                text = file_content.decode("utf-8")
            elif uploaded_file.type == "application/pdf":
                doc = fitz.open(stream=file_content, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(BytesIO(file_content))
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                st.warning(f"Unsupported file type: {uploaded_file.type}")
                continue

            docs.append(Document(text=text, metadata={"file_name": uploaded_file.name}))
        except Exception as e:
            st.error(f"Error loading {uploaded_file.name}: {e}")
            continue

    return docs
